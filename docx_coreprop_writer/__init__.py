#!/usr/bin/env python3
from typing import List, Dict

import datetime
import sys
import argparse
import yaml
from box import Box
import docx
from docx.section import Section
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.table import Table, _Cell
from docx.text.paragraph import Paragraph, Run
from docx.enum.section import WD_SECTION
from docx_coreprop_writer.version import version

META_KEY = "docx_coreprop"
ATTR_LIST = ["author",
             "category",
             "comments",
             "content_status",
             "created",
             "identifier",
             "keywords",
             "language",
             "last_modified_by",
             "last_printed",
             "modified",
             "revision",
             "subject",
             "title",
             "version",
             ]

TABLE_ALIGNMENT_IN_PAGE = {"left": WD_TABLE_ALIGNMENT.LEFT,
                           "center": WD_TABLE_ALIGNMENT.CENTER,
                           "right": WD_TABLE_ALIGNMENT.RIGHT}

CELL_VERTICAL_ALIGMENT = {"top": WD_CELL_VERTICAL_ALIGNMENT.TOP,
                          "center": WD_CELL_VERTICAL_ALIGNMENT.CENTER,
                          "bottom": WD_CELL_VERTICAL_ALIGNMENT.BOTTOM,
                          "both": WD_CELL_VERTICAL_ALIGNMENT.BOTH
                          }


def ensure_value(namespace, dest, default):
    """ Thanks to https://stackoverflow.com/a/29335524/6592473 """

    stored = getattr(namespace, dest, None)
    if stored is None:
        return default
    return stored


class StoreDict(argparse.Action):
    """ Thanks to https://stackoverflow.com/a/29335524/6592473 """

    def __call__(self, parser, namespace, values, option_string=None):
        vals = dict(ensure_value(namespace, self.dest, {}))
        key, _, val = values.partition('=')
        vals[key] = val
        setattr(namespace, self.dest, vals)


def get_choice(meta_ext: Dict or None, meta_file: Dict, key):
    """ tries to get meta_ext[key], then try meta_file[key]

    :param Dict or None meta_ext:
    :param Dict meta_file:
    :param str key:
    :return Dict ret:

    Uses metadata from external when exist (meta_ext), otherwise from default (meta_file)
    """
    assert meta_file is not None
    if meta_ext is not None:
        ret = meta_ext.get(key, None)
        if ret is None:
            ret = meta_file.get(key)
    else:
        ret = meta_file.get(key, None)
    return ret


class DictDotNotation(dict):
    def __init__(self, *args, **kwargs):
        super().__init__(*args, **kwargs)
        self.__dict__ = self


def apply_core_properties(meta_file, filename):
    """ Overwrite DOCX core property from meta_file or meta_ext dictionaries
    When both dict has value for each for same key, meta_ext has priority

    :param dict meta_file:
    :param str filename:
    """

    doc = docx.Document(filename)  # type:docx.Document

    meta = Box({key: meta_file.get(key) for key in ATTR_LIST})
    [print("{} = {}".format(key, val), file=sys.stderr) for key, val in meta.items()]
    if meta.author is not None:
        """ author (unicode)
        Note: named `creator` in spec.
        An entity primarily responsible for making the content of the resource. (Dublin Core)
        """
        doc.core_properties.author = meta.author
    if meta.category is not None:
        """ category (unicode)
        A categorization of the content of this package.
        Example values for this property might include: Resume, Letter, Financial Forecast, Proposal,
        Technical Presentation, and so on. (Open Packaging Conventions)
        """
        doc.core_properties.category = meta.category
    if meta.comments is not None:
        """comments (unicode)
        Note: named `description` in spec.
        An explanation of the content of the resource.
        Values might include an abstract, table of contents, reference to a graphical representation
        of content, and a free-text account of the content. (Dublin Core)
        """
        doc.core_properties.comments = meta.comments
    if meta.content_status is not None:
        """content_status (unicode)
        The status of the content.
        Values might include "Draft", "Reviewed", and "Final". (Open Packaging Conventions)
        """
        doc.core_properties.content_status = meta.content_status
    if meta.created is not None:
        """created (datetime)
        Date of creation of the resource. (Dublin Core)
        """
        doc.core_properties.created = datetime.datetime.strptime(meta.created, "%d-%b-%Y")  # DD-MMM-YYYY
    if meta.identifier is not None:
        """identifier (unicode)
        An unambiguous reference to the resource within a given context. (Dublin Core)
        """
        doc.core_properties.identifier = meta.identifier
    if meta.keywords is not None:
        """keywords (unicode)
        A delimited set of keywords to support searching and indexing.
        This is typically a list of terms that are not available elsewhere
        in the properties. (Open Packaging Conventions)
        """
        doc.core_properties.keywords = meta.keywords
    if meta.language is not None:
        """language (unicode)
        The language of the intellectual content of the resource. (Dublin Core)
        """
        doc.core_properties.language = meta.language
    if meta.last_modified_by is not None:
        """last_modified_by (unicode)
        The user who performed the last modification. The identification is environment-specific.
        Examples include a name, email address, or employee ID.
        It is recommended that this value be as concise as possible. (Open Packaging Conventions)
        """
        doc.core_properties.last_modified_by = meta.last_modified_by
    if meta.last_printed is not None:
        """last_printed (datetime)
        The date and time of the last printing. (Open Packaging Conventions)
        """
        doc.core_properties.last_printed = datetime.datetime.strptime(meta.last_printed, "%d-%b-%Y")
    if meta.modified is not None:
        """modified (datetime)
        Date on which the resource was changed. (Dublin Core)
        """
        doc.core_properties.modified = datetime.datetime.strptime(meta.modified, "%d-%b-%Y")
    if meta.revision is not None:
        """revision (int)
        The revision number. This value might indicate the number of saves or revisions,
        provided the application updates it after each revision. (Open Packaging Conventions)
        """
        doc.core_properties.revision = meta.revision
    if meta.subject is not None:
        """subject (unicode)
        The topic of the content of the resource. (Dublin Core)
        """
        doc.core_properties.subject = meta.subject
    if meta.title is not None:
        """title (unicode)
        The name given to the resource. (Dublin Core)
        """
        doc.core_properties.title = meta.title
    if meta.version is not None:
        """version (unicode)
        The version designator.
        This value is set by the user or by the application. (Open Packaging Conventions)
        """
        doc.core_properties.version = meta.version

    doc.save(filename)


def apply_table_alignment_in_page(meta_file, filename):
    """
    :param dict meta_file:
    :param str filename:
    :return:
    """
    _message = "Each table has aligned at {} of page"
    _key = "table-alignment-in-page"

    table_alignment_in_page = meta_file.get(_key)

    if table_alignment_in_page is not None:
        doc = docx.Document(filename)  # type:docx.Document
        table_alignment_in_page = table_alignment_in_page.lower()
        print(_message.format(table_alignment_in_page), file=sys.stderr)
        table: Table
        for table in doc.tables:
            table.alignment = TABLE_ALIGNMENT_IN_PAGE[table_alignment_in_page]
        doc.save(filename)


def apply_cell_vertical_alignment(meta_file, filename):
    """
    :param dict meta_file:
    :param str filename:
    :return:
    """
    _message = "Each table cell has vertically {} aligned"
    _key = "table-cell-vertical-alignment"

    cell_vertical_alignment = meta_file.get(_key)

    if cell_vertical_alignment is not None:
        doc = docx.Document(filename)  # type:docx.Document
        cell_vertical_alignment = cell_vertical_alignment.lower()
        print(_message.format(cell_vertical_alignment), file=sys.stderr)
        table: Table
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    cell.vertical_alignment = CELL_VERTICAL_ALIGMENT[cell_vertical_alignment]
        doc.save(filename)


def unset_word2010_compatibility_mode(meta_file, filename):
    """
    :param dict meta_file:
    :param str filename:
    :return:
    """
    _message = "Drop Word 2010 compatibility mode"
    _key = "word2010compatible"
    w_compat = "w:compat"
    subelements = [
        ("w:useFELayout", {}),
        ("w:compatSetting", {qn("w:name"): "compatibilityMode",
                             qn("w:url"): "http://schemas.microsoft.com/office/word",
                             qn("w:val"): "15"
                             },
         ),
        ("w:compatSetting", {qn("w:name"): "overrideTableStyleFontSizeAndJustification",
                             qn("w:url"): "http://schemas.microsoft.com/office/word",
                             qn("w:val"): "1",
                             },
         ),
        ("w:compatSetting", {qn("w:name"): "enableOpenTypeFeatures",
                             qn("w:url"): "http://schemas.microsoft.com/office/word",
                             qn("w:val"): "1"
                             },
         ),
        ("w:compatSetting", {qn("w:name"): "doNotFlipMirrorIndents",
                             qn("w:url"): "http://schemas.microsoft.com/office/word",
                             qn("w:val"): "1",
                             },
         ),
        ("w:compatSetting", {qn("w:name"): "differentiateMultirowTableHeaders",
                             qn("w:url"): "http://schemas.microsoft.com/office/word",
                             qn("w:val"): "1",
                             },
         ),
        ("w:compatSetting", {qn("w:name"): "useWord2013TrackBottomHyphenation",
                             qn("w:url"): "http://schemas.microsoft.com/office/word",
                             qn("w:val"): "0",
                             },
         ),
    ]

    word2010compatible = meta_file.get(_key, False)

    if word2010compatible is False:
        print(_message, file=sys.stderr)
        doc = docx.Document(filename)  # type:docx.Document
        doc.settings.element.remove_all(w_compat)

        compat = OxmlElement(w_compat)
        for sub_elem in subelements:
            subelement = OxmlElement(sub_elem[0], attrs=sub_elem[1])
            compat.append(subelement)
        doc.settings.element.append(compat)
        doc.save(filename)


def disable_table_autofit(meta_file, filename):
    """
    :param dict meta_file:
    :param str filename:
    :return:
    """
    _message = "Fix table column widths"
    _key = "disable-table-autofit"

    disable_table_autofit_meta = meta_file.get(_key, False)

    if disable_table_autofit_meta is True:
        doc = docx.Document(filename)  # type:docx.Document
        print(_message, file=sys.stderr)
        table: Table
        for table in doc.tables:
            table.autofit = False
        doc.save(filename)


def recommend_readonly(meta_file, filename):
    """
    :param dict meta_file:
    :param str filename:
    :return:
    """
    _message = "Set read only recommend flag"
    _key = "read-only-recommended"
    elem_name = "w:writeProtection"
    attr_name = qn("w:recommended")

    read_only = meta_file.get(_key, False)

    if read_only is True:
        print(_message, file=sys.stderr)
        doc = docx.Document(filename)  # type:docx.Document
        write_protection = doc.settings.element.xpath(elem_name)
        if write_protection == []:
            write_protection = OxmlElement(elem_name, attrs={attr_name: "1"})
            doc.settings.element.append(write_protection)
        else:
            write_protection = write_protection[0]
            if write_protection.get(attr_name, None) is None:
                write_protection.set(attr_name, "1")

        doc.save(filename)


def replace_table_style(meta_file, filename):
    """
    :param dict meta_file:
    :param str filename:
    :return:
    """
    _message = "Replace table styles"
    _key = "table"

    table = meta_file.get(_key)

    if table is not None:
        print(_message, file=sys.stderr)

        doc = docx.Document(filename)  # type:docx.Document
        for key, val in table.items():
            for table in doc.tables:
                if table.style.name == key:
                    print("{} -> {}".format(key, val), file=sys.stderr)
                    table.style = doc.styles[val]
        doc.save(filename)


def replace_paragraph_style(meta_file, filename):
    """
    :param dict meta_file:
    :param str filename:
    :return:
    """
    _message = "Replace paragraph styles"
    _key = "paragraph"

    para = meta_file.get(_key)

    if para is not None:
        print(_message, file=sys.stderr)

        doc = docx.Document(filename)  # type:docx.Document
        for key, val in para.items():
            for para in doc.paragraphs:
                if para.style.name == key:
                    print("{} -> {}".format(key, val), file=sys.stderr)
                    para.style = doc.styles[val]

        doc.save(filename)


def replace_character_style(meta_file, filename):
    """
    :param dict meta_file:
    :param str filename:
    :return:
    """
    _message = "Replace character styles"
    _key = "character"

    char = meta_file.get(_key)

    if char is not None:
        print(_message, file=sys.stderr)
        doc = docx.Document(filename)  # type:docx.Document
        for key, val in char.items():
            para: Paragraph
            for para in doc.paragraphs:
                run: Run
                for run in para.runs:
                    if run.style.name == key:
                        print("{} -> {}".format(key, val), file=sys.stderr)
                        run.style = doc.styles[val]

        doc.save(filename)


def insert_extra_section(meta_file, filename):
    """
    :param dict meta_file:
    :param str filename:
    :return:
    """
    _message = "Insert extra section (clears Header/Footer content)"
    _key = "extra_section"

    char = meta_file.get(_key, False)

    if char is True:
        print(_message, file=sys.stderr)
        doc: docx.Document = docx.Document(filename)

        last_section: Section = doc.sections[-1]
        extra_section: Section = doc.add_section(WD_SECTION.NEW_PAGE)

        extra_section.orientation = last_section.orientation

        extra_section.page_width = last_section.page_width
        extra_section.page_height = last_section.page_height

        extra_section.left_margin = last_section.left_margin
        extra_section.right_margin = last_section.right_margin
        extra_section.top_margin = last_section.top_margin
        extra_section.bottom_margin = last_section.bottom_margin

        extra_section.different_first_page_header_footer = True
        extra_section.first_page_header.is_linked_to_previous = False
        extra_section.first_page_footer.is_linked_to_previous = False
        extra_section.header.is_linked_to_previous = False
        extra_section.footer.is_linked_to_previous = False
        extra_section.even_page_header.is_linked_to_previous = False
        extra_section.even_page_footer.is_linked_to_previous = False

        doc.save(filename)


def insert_okuzuke_table(meta_file, filename):
    """
    :param dict meta_file:
    :param str filename:
    :return:
    """
    _message = "Insert Okuzuke table"
    _key = "okuzuke"

    okuzuke = meta_file.get(_key)

    if okuzuke is not None:
        print(_message, file=sys.stderr)
        doc: docx.Document = docx.Document(filename)
        last_section: Section = doc.sections[-1]

        vAlign = OxmlElement("w:vAlign", attrs={qn("w:val"): "bottom"})
        last_section._sectPr.append(vAlign)

        doc.add_page_break()

        table: Table = doc.add_table(rows=0, cols=1, style=okuzuke.get("table-style", "Normal Table"))
        table.alignment = TABLE_ALIGNMENT_IN_PAGE["center"]
        for row_text in okuzuke.get("rows", []):
            row = table.add_row()
            row.cells[0].text = row_text.strip()
            row.cells[0].paragraphs[0].style = okuzuke.get("para-style", "Table Body Left")

        doc.add_page_break()
        doc.add_page_break()

        doc.save(filename)


def main():
    parser = argparse.ArgumentParser(description="Reads yaml, overwrites DOCX core property")
    parser.add_argument("--input", "-I", required=True, default=None, help="yaml input filename")
    parser.add_argument("--output", "-O", required=True, help="docx output filename")
    parser.add_argument("--metadata", "-M", default={}, action=StoreDict)
    # parser.add_argument("--paragraph", "-P", default=None, action=StoreDict)
    # parser.add_argument("--table", "-T", default=None, action=StoreDict)
    parser.add_argument('--version', action='version', version=str(version))

    args = parser.parse_args()

    meta_file = Box.from_yaml(filename=args.input).get(META_KEY, Box({}))
    doc = args.output
    meta_ext = Box(args.metadata)
    # style_ext = {"paragraph": args.paragraph, "table": args.table, }
    metadata = meta_file + meta_ext

    unset_word2010_compatibility_mode(metadata, doc)
    apply_core_properties(metadata, doc)
    replace_paragraph_style(metadata, doc)
    insert_extra_section(metadata, doc)
    replace_table_style(metadata, doc)
    replace_character_style(metadata, doc)
    apply_table_alignment_in_page(metadata, doc)
    apply_cell_vertical_alignment(metadata, doc)
    disable_table_autofit(metadata, doc)
    recommend_readonly(metadata, doc)
    insert_okuzuke_table(metadata, doc)

    print("{} processed".format(doc), file=sys.stderr)


if __name__ == "__main__":
    main()
